# proposal_generator.py
import os
import uuid
from datetime import datetime
from typing import Dict, List, Any
import anthropic
from real_analysis_system import get_real_analysis_results

class ProposalGenerator:
    """Advanced proposal generation engine with multiple deliverable types"""

    def __init__(self, project, analysis_results, company_name="ITSS Global", contact_person="Project Manager"):
        self.project = project
        self.analysis_results = analysis_results
        self.company_name = company_name
        self.contact_person = contact_person
        self.client_name = getattr(project, 'client_name', 'Valued Client') or 'Valued Client'

        # Initialize AI client
        self.anthropic_key = os.getenv('ANTHROPIC_API_KEY')
        if self.anthropic_key:
            self.client = anthropic.Anthropic(api_key=self.anthropic_key)
        else:
            self.client = None

        # Ensure output directory exists
        self.output_dir = 'generated_proposals'
        os.makedirs(self.output_dir, exist_ok=True)

    def generate_document(self, deliverable_type: str, output_format: str = 'pdf', detail_level: str = 'standard') -> Dict[str, Any]:
        """Generate a specific type of document"""

        # Get the appropriate content generator
        content_generators = {
            'technical': self._generate_technical_proposal,
            'commercial': self._generate_commercial_proposal,
            'implementation': self._generate_implementation_plan,
            'architecture': self._generate_architecture_document,
            'company': self._generate_company_profile,
            'compliance': self._generate_compliance_matrix
        }

        if deliverable_type not in content_generators:
            raise ValueError(f"Unknown deliverable type: {deliverable_type}")

        # Generate content using AI
        content = content_generators[deliverable_type](detail_level)

        # Create document metadata
        doc_info = {
            'type': deliverable_type,
            'title': self._get_document_title(deliverable_type),
            'description': self._get_document_description(deliverable_type),
            'format': output_format,
            'detail_level': detail_level,
            'generated_at': datetime.now().isoformat(),
            'content': content
        }

        # Save document
        filename = self._save_document(doc_info, output_format)

        # Return download info
        return {
            'title': doc_info['title'],
            'description': doc_info['description'],
            'filename': filename,
            'format': output_format,
            'size': self._get_file_size(filename),
            'download_url': f'/download-proposal/{filename}',
            'filepath': os.path.join(self.output_dir, filename),
            'generated_at': doc_info['generated_at']
        }

    def generate_custom_document(self, custom_deliverable, output_format: str = 'pdf', detail_level: str = 'standard') -> Dict[str, Any]:
        """Generate a document based on custom deliverable template"""
        
        if not self.client:
            content = self._fallback_custom_document(custom_deliverable)
        else:
            # Replace placeholders in the prompt template
            prompt_template = custom_deliverable.prompt_template or ""
            
            # Available placeholders
            placeholders = {
                'project_name': self.project.name,
                'client_name': self.client_name,
                'company_name': self.company_name,
                'contact_person': self.contact_person,
                'requirements': self._format_requirements(),
                'analysis_summary': self._format_analysis_summary(),
                'key_constraints': self._format_constraints(),
                'project_details': self._format_project_details()
            }
            
            # Replace placeholders in template
            formatted_prompt = prompt_template
            for placeholder, value in placeholders.items():
                formatted_prompt = formatted_prompt.replace('{' + placeholder + '}', str(value))
            
            # Add detail level instructions
            detail_instructions = {
                'executive': "Keep the response concise and executive-focused (5-10 pages).",
                'standard': "Provide a comprehensive response with good detail (10-15 pages).",
                'comprehensive': "Provide an extremely detailed and thorough response (15+ pages)."
            }
            
            full_prompt = f"""
            {formatted_prompt}
            
            **Instructions:**
            - {detail_instructions.get(detail_level, detail_instructions['standard'])}
            - Format as professional business document with clear sections
            - Use markdown formatting for structure
            - Include specific examples and recommendations where applicable
            - If the tender / RFP document insists of submitting the proposal in a specific format, then include those format / content.
            - Ensure all content is relevant to the project context provided
            """
            
            try:
                response = self.client.messages.create(
                    model="claude-sonnet-4-20250514",
                    max_tokens=4000,
                    messages=[{"role": "user", "content": full_prompt}]
                )
                content = response.content[0].text
            except Exception as e:
                print(f"AI generation failed: {e}")
                content = self._fallback_custom_document(custom_deliverable)

        # Create document metadata
        doc_info = {
            'title': custom_deliverable.title,
            'description': custom_deliverable.description,
            'type': f'custom_{custom_deliverable.id}',
            'content': content,
            'generated_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'project_id': self.project.id,
            'project_name': self.project.name,
            'company_name': self.company_name,
            'contact_person': self.contact_person
        }

        # Save document  
        filename = self._save_document(doc_info, output_format)

        # Return download info
        return {
            'title': doc_info['title'],
            'description': doc_info['description'],
            'filename': filename,
            'format': output_format,
            'size': self._get_file_size(filename),
            'download_url': f'/download-proposal/{filename}',
            'filepath': os.path.join(self.output_dir, filename),
            'generated_at': doc_info['generated_at']
        }

    def _generate_technical_proposal(self, detail_level: str) -> str:
        """Generate comprehensive technical proposal with past proposal intelligence"""

        if not self.client:
            return self._fallback_technical_proposal()

        # Get relevant past proposals for technical content
        past_proposal_insights = self._get_past_proposal_insights("technical")
        
        # Get relevant partner solutions
        partner_solutions = self._get_partner_solutions()
        
        # Build enhanced prompt with past proposal content and partner solutions
        prompt = f"""
        You are a Senior Technical Architect creating a comprehensive technical proposal for:

        **PROJECT:** {self.project.name}
        **CLIENT:** {self.client_name}
        **BIDDER:** {self.company_name} (Certified Temenos Implementation Partner specializing in BFSI solutions)

        **RFP ANALYSIS RESULTS:**

        **Must-Have Requirements:**
        {chr(10).join('• ' + req for req in self.analysis_results.get('must_have_requirements', []))}

        **Technical Specifications:**
        {chr(10).join('• ' + spec for spec in self.analysis_results.get('technical_specifications', []))}

        **Project Details:**
        - Timeline: {self.analysis_results.get('project_details', {}).get('timeline', 'To be determined')}
        - Budget: {self.analysis_results.get('project_details', {}).get('budget', 'To be determined')}
        - Evaluation: {self.analysis_results.get('project_details', {}).get('evaluation_criteria', 'Standard criteria')}

        {self._format_past_proposal_context(past_proposal_insights)}

        {self._format_partner_solutions_context(partner_solutions)}

        Generate a comprehensive technical proposal with the following structure:

        # TECHNICAL PROPOSAL
        ## For {self.project.name}

        ### 1. EXECUTIVE SUMMARY
        [Compelling 2-3 paragraph overview of our technical solution]

        ### 2. UNDERSTANDING OF REQUIREMENTS
        [Demonstrate deep understanding of client needs and technical challenges]

        ### 3. PROPOSED SOLUTION ARCHITECTURE
        #### 3.1 High-Level Architecture
        [Overall system architecture approach]

        #### 3.2 Technology Stack
        [Recommended technologies with justifications]

        #### 3.3 Integration Strategy
        [How system integrates with existing infrastructure]

        ### 4. TECHNICAL IMPLEMENTATION APPROACH
        #### 4.1 Development Methodology
        [Agile/DevOps approach]

        #### 4.2 Quality Assurance Framework
        [Testing and QA processes]

        #### 4.3 Security Implementation
        [Comprehensive security measures]

        ### 5. SCALABILITY & PERFORMANCE
        [How solution handles growth and performance requirements]

        ### 6. COMPLIANCE & STANDARDS
        [Adherence to industry standards and regulations]

        ### 7. RISK MITIGATION
        [Technical risks and mitigation strategies]

        ### 8. SUPPORT & MAINTENANCE
        [Ongoing support model and maintenance approach]

        **IMPORTANT GUIDELINES:**
        - Write in professional, confident tone
        - Include specific technical details and justifications
        - Address ALL must-have requirements explicitly
        - Use industry best practices and current technologies
        - Detail level: {detail_level}
        - Make it comprehensive yet readable
        - Include realistic implementation approaches

        Generate a detailed, professional technical proposal that would win this RFP.
        """

        try:
            response = self.client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=4000,
                messages=[{"role": "user", "content": prompt}]
            )
            return response.content[0].text
        except Exception as e:
            print(f"AI generation failed: {e}")
            return self._fallback_technical_proposal()

    def _generate_commercial_proposal(self, detail_level: str) -> str:
        """Generate comprehensive commercial proposal with past proposal intelligence"""

        if not self.client:
            return self._fallback_commercial_proposal()

        # Get relevant past commercial proposals
        past_proposal_insights = self._get_past_proposal_insights("commercial")
        
        # Get relevant partner solutions for commercial considerations
        partner_solutions = self._get_partner_solutions()

        prompt = f"""
        You are a Commercial Director creating a compelling commercial proposal for:

        **PROJECT:** {self.project.name}
        **CLIENT:** {self.client_name}
        **BIDDER:** {self.company_name} (Certified Temenos Implementation Partner specializing in BFSI solutions)

        **PROJECT REQUIREMENTS:**
        Must-Have Requirements: {len(self.analysis_results.get('must_have_requirements', []))} items
        Good-to-Have Requirements: {len(self.analysis_results.get('good_to_have_requirements', []))} items
        Technical Specifications: {len(self.analysis_results.get('technical_specifications', []))} items

        **PROJECT DETAILS:**
        - Timeline: {self.analysis_results.get('project_details', {}).get('timeline', 'To be determined')}
        - Budget Context: {self.analysis_results.get('project_details', {}).get('budget', 'Competitive pricing required')}

        {self._format_past_proposal_context(past_proposal_insights)}

        {self._format_partner_solutions_context(partner_solutions)}

        Generate a comprehensive commercial proposal with this structure:

        # COMMERCIAL PROPOSAL
        ## For {self.project.name}

        ### 1. EXECUTIVE SUMMARY
        [Strong value proposition and commercial overview]

        ### 2. PRICING STRUCTURE
        #### 2.1 Cost Breakdown
        [Detailed cost breakdown with phases]

        #### 2.2 Payment Schedule
        [Milestone-based payment terms]

        #### 2.3 Pricing Model
        [Fixed price, time & materials, or hybrid approach]

        ### 3. PROJECT INVESTMENT & ROI
        #### 3.1 Total Cost of Ownership
        [5-year TCO analysis]

        #### 3.2 Return on Investment
        [Expected ROI and payback period]

        #### 3.3 Value Proposition
        [Business benefits and value delivered]

        ### 4. PROJECT TIMELINE & RESOURCES
        #### 4.1 Project Phases
        [Implementation timeline with milestones]

        #### 4.2 Team Structure
        [Resource allocation and team composition]

        #### 4.3 Resource Requirements
        [Client resource commitments needed]

        ### 5. COMMERCIAL TERMS
        #### 5.1 Contract Terms
        [Key commercial conditions]

        #### 5.2 Service Level Agreements
        [Performance guarantees and SLAs]

        #### 5.3 Support & Maintenance
        [Ongoing support pricing and terms]

        ### 6. RISK ASSESSMENT
        [Commercial risks and mitigation strategies]

        ### 7. WHY CHOOSE {self.company_name.upper()}
        [Competitive advantages and differentiators]

        **PRICING GUIDELINES:**
        - Include realistic market-based pricing
        - Show cost breakdown for transparency
        - Justify pricing with value delivered
        - Include options for different budget levels
        - Detail level: {detail_level}

        Generate a compelling commercial proposal that demonstrates value and wins business.
        """

        try:
            response = self.client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=4000,
                messages=[{"role": "user", "content": prompt}]
            )
            return response.content[0].text
        except Exception as e:
            print(f"AI generation failed: {e}")
            return self._fallback_commercial_proposal()

    def _generate_implementation_plan(self, detail_level: str) -> str:
        """Generate detailed implementation plan"""

        if not self.client:
            return self._fallback_implementation_plan()

        prompt = f"""
        You are a Senior Project Manager creating a detailed implementation plan for:

        **PROJECT:** {self.project.name}
        **CLIENT:** {self.client_name}
        **SCOPE:** Implementation of solution meeting {len(self.analysis_results.get('must_have_requirements', []))} must-have requirements

        **KEY REQUIREMENTS TO IMPLEMENT:**
        {chr(10).join('• ' + req for req in self.analysis_results.get('must_have_requirements', [])[:10])}

        **PROJECT CONTEXT:**
        - Timeline: {self.analysis_results.get('project_details', {}).get('timeline', 'Standard project timeline')}
        - Complexity: {len(self.analysis_results.get('technical_specifications', []))} technical specifications

        Generate a comprehensive implementation plan:

        # IMPLEMENTATION PLAN
        ## For {self.project.name}

        ### 1. PROJECT OVERVIEW
        [Project objectives, scope, and success criteria]

        ### 2. WORK BREAKDOWN STRUCTURE (WBS)
        #### Phase 1: Project Initiation & Planning (Weeks 1-2)
        [Detailed tasks and deliverables]

        #### Phase 2: System Design & Architecture (Weeks 3-6)
        [Design phase tasks and milestones]

        #### Phase 3: Development & Configuration (Weeks 7-16)
        [Core development activities]

        #### Phase 4: Integration & Testing (Weeks 17-20)
        [Testing and integration activities]

        #### Phase 5: Deployment & Go-Live (Weeks 21-24)
        [Deployment and launch activities]

        #### Phase 6: Support & Optimization (Weeks 25-26)
        [Post-launch support and optimization]

        ### 3. PROJECT TIMELINE
        [Gantt chart style timeline with dependencies]

        ### 4. RESOURCE ALLOCATION
        #### 4.1 Team Structure
        [Project team roles and responsibilities]

        #### 4.2 Resource Loading
        [Resource allocation by phase and skill type]

        #### 4.3 Client Resource Requirements
        [Required client involvement and resources]

        ### 5. RISK MANAGEMENT
        #### 5.1 Risk Register
        [Key project risks and impact assessment]

        #### 5.2 Mitigation Strategies
        [Risk mitigation and contingency plans]

        #### 5.3 Issue Escalation Process
        [How issues will be handled and escalated]

        ### 6. QUALITY ASSURANCE
        #### 6.1 QA Framework
        [Quality assurance processes and checkpoints]

        #### 6.2 Testing Strategy
        [Comprehensive testing approach]

        #### 6.3 Acceptance Criteria
        [Clear acceptance criteria for each phase]

        ### 7. COMMUNICATION PLAN
        [Stakeholder communication and reporting structure]

        ### 8. CHANGE MANAGEMENT
        [How changes will be managed and controlled]

        **GUIDELINES:**
        - Include realistic timelines and resource estimates
        - Address all major requirements identified in the RFP
        - Show clear milestones and deliverables
        - Detail level: {detail_level}
        - Include contingencies and risk mitigation

        Create a professional, executable implementation plan.
        """

        try:
            response = self.client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=4000,
                messages=[{"role": "user", "content": prompt}]
            )
            return response.content[0].text
        except Exception as e:
            print(f"AI generation failed: {e}")
            return self._fallback_implementation_plan()

    def _generate_architecture_document(self, detail_level: str) -> str:
        """Generate technical architecture document"""

        if not self.client:
            return self._fallback_architecture_document()

        prompt = f"""
        You are a Lead Solutions Architect creating a detailed technical architecture document for:

        **PROJECT:** {self.project.name}
        **TECHNICAL REQUIREMENTS:**
        {chr(10).join('• ' + spec for spec in self.analysis_results.get('technical_specifications', []))}

        **SYSTEM REQUIREMENTS:**
        Must-Have: {len(self.analysis_results.get('must_have_requirements', []))} requirements
        Good-to-Have: {len(self.analysis_results.get('good_to_have_requirements', []))} requirements

        Generate a comprehensive technical architecture document:

        # TECHNICAL ARCHITECTURE DOCUMENT
        ## For {self.project.name}

        ### 1. ARCHITECTURE OVERVIEW
        [High-level architecture vision and principles]

        ### 2. SYSTEM ARCHITECTURE
        #### 2.1 Logical Architecture
        [Logical system components and their relationships]

        #### 2.2 Physical Architecture
        [Physical deployment and infrastructure]

        #### 2.3 Technology Stack
        [Detailed technology choices with justifications]

        ### 3. COMPONENT DESIGN
        #### 3.1 Application Layer
        [Application components and services]

        #### 3.2 Data Layer
        [Database design and data management]

        #### 3.3 Integration Layer
        [APIs, messaging, and integration patterns]

        ### 4. SECURITY ARCHITECTURE
        #### 4.1 Security Framework
        [Comprehensive security design]

        #### 4.2 Authentication & Authorization
        [Identity and access management]

        #### 4.3 Data Security
        [Data protection and encryption]

        ### 5. PERFORMANCE & SCALABILITY
        #### 5.1 Performance Requirements
        [Performance targets and optimization]

        #### 5.2 Scalability Design
        [How system scales with growth]

        #### 5.3 Load Balancing & Clustering
        [High availability design]

        ### 6. DATA ARCHITECTURE
        #### 6.1 Data Model
        [Conceptual and logical data models]

        #### 6.2 Data Flow
        [Data movement and transformation]

        #### 6.3 Data Governance
        [Data quality and governance]

        ### 7. INFRASTRUCTURE REQUIREMENTS
        #### 7.1 Hardware Specifications
        [Server and infrastructure requirements]

        #### 7.2 Network Architecture
        [Network design and requirements]

        #### 7.3 Cloud Strategy
        [Cloud deployment strategy if applicable]

        ### 8. DISASTER RECOVERY & BACKUP
        [Business continuity and disaster recovery]

        **TECHNICAL GUIDELINES:**
        - Use current industry best practices
        - Include specific technology recommendations
        - Address scalability and performance requirements
        - Detail level: {detail_level}
        - Ensure enterprise-grade architecture

        Create a detailed, implementable technical architecture.
        """

        try:
            response = self.client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=4000,
                messages=[{"role": "user", "content": prompt}]
            )
            return response.content[0].text
        except Exception as e:
            print(f"AI generation failed: {e}")
            return self._fallback_architecture_document()

    def _generate_company_profile(self, detail_level: str) -> str:
        """Generate company profile and credentials"""

        if not self.client:
            return self._fallback_company_profile()

        prompt = f"""
        You are a Business Development Director creating a compelling company profile for:

        **BIDDING COMPANY:** {self.company_name}
        **PROJECT:** {self.project.name}
        **CLIENT:** {self.client_name}
        **CONTACT:** {self.contact_person}

        **PROJECT CONTEXT:**
        This is for a {self.project.description or 'technology implementation project'} requiring expertise in:
        {chr(10).join('• ' + spec for spec in self.analysis_results.get('technical_specifications', [])[:5])}

        Generate a professional company profile:

        # COMPANY PROFILE & CREDENTIALS
        ## {self.company_name}

        ### 1. COMPANY OVERVIEW
        #### 1.1 About {self.company_name}
        [Compelling company introduction and mission]

        #### 1.2 Company History
        [Brief history and key milestones]

        #### 1.3 Vision & Values
        [Company vision, mission, and core values]

        ### 2. CORE COMPETENCIES
        #### 2.1 Technical Expertise
        [Relevant technical capabilities and specializations]

        #### 2.2 Industry Experience
        [Industry knowledge and sector expertise]

        #### 2.3 Service Offerings
        [Complete service portfolio]

        ### 3. RELEVANT PROJECT EXPERIENCE
        #### 3.1 Similar Projects
        [Case studies of similar successful projects]

        #### 3.2 Client Success Stories
        [Specific examples of client value delivered]

        #### 3.3 Key Achievements
        [Notable project outcomes and metrics]

        ### 4. TEAM CREDENTIALS
        #### 4.1 Leadership Team
        [Key executives and their backgrounds]

        #### 4.2 Technical Team
        [Technical team capabilities and certifications]

        #### 4.3 Project Team for This Engagement
        [Proposed team structure and key personnel]

        ### 5. CERTIFICATIONS & PARTNERSHIPS
        #### 5.1 Industry Certifications
        [Relevant certifications and accreditations]

        #### 5.2 Technology Partnerships
        [Strategic technology partnerships]

        #### 5.3 Awards & Recognition
        [Industry awards and recognition]

        ### 6. FINANCIAL STABILITY
        #### 6.1 Company Financial Health
        [Financial stability and growth]

        #### 6.2 Insurance & Bonding
        [Insurance coverage and bonding capacity]

        ### 7. CLIENT REFERENCES
        [Relevant client testimonials and references]

        ### 8. WHY CHOOSE {self.company_name.upper()}
        [Unique value proposition and differentiators]

        **GUIDELINES:**
        - Make it compelling and professional
        - Focus on relevant experience and capabilities
        - Include specific achievements and metrics
        - Detail level: {detail_level}
        - Tailor to the project requirements

        Create a persuasive company profile that builds confidence.
        """

        try:
            response = self.client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=4000,
                messages=[{"role": "user", "content": prompt}]
            )
            return response.content[0].text
        except Exception as e:
            print(f"AI generation failed: {e}")
            return self._fallback_company_profile()

    def _generate_compliance_matrix(self, detail_level: str) -> str:
        """Generate compliance and requirements matrix"""

        if not self.client:
            return self._fallback_compliance_matrix()

        prompt = f"""
        You are a Compliance Manager creating a detailed compliance matrix for:

        **PROJECT:** {self.project.name}
        **REQUIREMENTS TO ADDRESS:**

        **Must-Have Requirements:**
        {chr(10).join(f'{i+1}. {req}' for i, req in enumerate(self.analysis_results.get('must_have_requirements', [])))}

        **Compliance Requirements:**
        {chr(10).join('• ' + req for req in self.analysis_results.get('compliance_requirements', []))}

        **Technical Specifications:**
        {chr(10).join('• ' + spec for spec in self.analysis_results.get('technical_specifications', []))}

        Generate a comprehensive compliance matrix:

        # COMPLIANCE & REQUIREMENTS MATRIX
        ## For {self.project.name}

        ### 1. EXECUTIVE SUMMARY
        [Overview of compliance approach and commitment]

        ### 2. REQUIREMENTS TRACEABILITY MATRIX

        | Requirement ID | Requirement Description | Compliance Status | Solution Approach | Validation Method |
        |----------------|------------------------|-------------------|-------------------|-------------------|
        [Create detailed table for each must-have requirement]

        ### 3. COMPLIANCE CHECKLIST
        #### 3.1 Functional Requirements Compliance
        [Detailed compliance for functional requirements]

        #### 3.2 Technical Requirements Compliance
        [Technical specifications compliance]

        #### 3.3 Regulatory Compliance
        [Industry and regulatory compliance]

        ### 4. STANDARDS & CERTIFICATIONS
        #### 4.1 Industry Standards
        [Relevant industry standards compliance]

        #### 4.2 Security Standards
        [Security and data protection compliance]

        #### 4.3 Quality Standards
        [Quality management and processes]

        ### 5. GAP ANALYSIS
        #### 5.1 Requirements Gaps (if any)
        [Any requirements not fully met and mitigation]

        #### 5.2 Enhancement Opportunities
        [Additional value beyond requirements]

        ### 6. TESTING & VALIDATION
        #### 6.1 Compliance Testing Strategy
        [How compliance will be tested and validated]

        #### 6.2 Acceptance Testing
        [User acceptance testing approach]

        #### 6.3 Documentation Requirements
        [Required documentation and deliverables]

        ### 7. RISK ASSESSMENT
        #### 7.1 Compliance Risks
        [Risks related to compliance and mitigation]

        #### 7.2 Regulatory Risks
        [Regulatory compliance risks]

        ### 8. ONGOING COMPLIANCE
        [How compliance will be maintained post-implementation]

        **COMPLIANCE GUIDELINES:**
        - Address every single requirement explicitly
        - Show clear traceability from requirement to solution
        - Include validation and testing methods
        - Detail level: {detail_level}
        - Demonstrate thorough understanding

        Create a comprehensive compliance matrix that proves full requirement coverage.
        """

        try:
            response = self.client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=4000,
                messages=[{"role": "user", "content": prompt}]
            )
            return response.content[0].text
        except Exception as e:
            print(f"AI generation failed: {e}")
            return self._fallback_compliance_matrix()

    def _save_document(self, doc_info: Dict[str, Any], output_format: str) -> str:
        """Save document in specified format"""

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{self.project.name}_{doc_info['type']}_{timestamp}.{output_format}"
        filename = filename.replace(' ', '_').replace('/', '_')

        filepath = os.path.join(self.output_dir, filename)

        if output_format.lower() == 'pdf':
            # For PDF generation, you'd typically use a library like reportlab or weasyprint
            # For now, save as HTML and note that PDF conversion is needed
            self._save_as_html(doc_info['content'], filepath.replace('.pdf', '.html'))
            filename = filename.replace('.pdf', '.html')

        elif output_format.lower() == 'docx':
            # For DOCX, you'd use python-docx
            self._save_as_docx(doc_info['content'], filepath)

        elif output_format.lower() == 'html':
            self._save_as_html(doc_info['content'], filepath)

        else:  # markdown or other text formats
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(doc_info['content'])

        return filename

    def _save_as_html(self, content: str, filepath: str) -> None:
        """Save content as formatted HTML"""
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>{self.project.name} - Proposal Document</title>
            <style>
                body {{ font-family: Arial, sans-serif; line-height: 1.6; margin: 40px; }}
                h1 {{ color: #2c3e50; border-bottom: 3px solid #3498db; padding-bottom: 10px; }}
                h2 {{ color: #34495e; border-bottom: 1px solid #bdc3c7; padding-bottom: 5px; }}
                h3 {{ color: #2c3e50; }}
                table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
                th, td {{ border: 1px solid #ddd; padding: 12px; text-align: left; }}
                th {{ background-color: #f2f2f2; }}
                .header {{ text-align: center; margin-bottom: 30px; }}
                .footer {{ margin-top: 50px; text-align: center; color: #7f8c8d; }}
                ul {{ margin: 10px 0; }}
                li {{ margin: 5px 0; }}
            </style>
        </head>
        <body>
            <div class="header">
                <h1>{self.project.name}</h1>
                <p>Prepared by: {self.company_name}</p>
                <p>Date: {datetime.now().strftime('%B %d, %Y')}</p>
            </div>

            {self._convert_markdown_to_html(content)}

            <div class="footer">
                <p>Generated by {self.company_name} | {datetime.now().strftime('%Y')}</p>
                <p>Contact: {self.contact_person}</p>
            </div>
        </body>
        </html>
        """

        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html_content)

    def _save_as_docx(self, content: str, filepath: str) -> None:
        """Save content as Word document"""
        try:
            from docx import Document
            from docx.shared import Inches

            doc = Document()

            # Add title
            title = doc.add_heading(f'{self.project.name} - Proposal', 0)
            title.alignment = 1  # Center alignment

            # Add metadata
            doc.add_paragraph(f'Prepared by: {self.company_name}')
            doc.add_paragraph(f'Contact: {self.contact_person}')
            doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')
            doc.add_page_break()

            # Add content (basic conversion from markdown-style to docx)
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('#### '):
                    doc.add_heading(line[5:], level=4)
                elif line.startswith('- ') or line.startswith('• '):
                    p = doc.add_paragraph(line[2:], style='List Bullet')
                elif line:
                    doc.add_paragraph(line)

            doc.save(filepath)

        except ImportError:
            # Fallback to text if python-docx not available
            with open(filepath.replace('.docx', '.txt'), 'w', encoding='utf-8') as f:
                f.write(content)

    def _convert_markdown_to_html(self, content: str) -> str:
        """Basic markdown to HTML conversion"""
        lines = content.split('\n')
        html_lines = []

        for line in lines:
            line = line.strip()
            if line.startswith('# '):
                html_lines.append(f'<h1>{line[2:]}</h1>')
            elif line.startswith('## '):
                html_lines.append(f'<h2>{line[3:]}</h2>')
            elif line.startswith('### '):
                html_lines.append(f'<h3>{line[4:]}</h3>')
            elif line.startswith('#### '):
                html_lines.append(f'<h4>{line[5:]}</h4>')
            elif line.startswith('- ') or line.startswith('• '):
                html_lines.append(f'<li>{line[2:]}</li>')
            elif line.startswith('**') and line.endswith('**'):
                html_lines.append(f'<p><strong>{line[2:-2]}</strong></p>')
            elif line:
                html_lines.append(f'<p>{line}</p>')
            else:
                html_lines.append('<br>')

        return '\n'.join(html_lines)

    def _get_document_title(self, deliverable_type: str) -> str:
        """Get human-readable document title"""
        titles = {
            'technical': 'Technical Proposal',
            'commercial': 'Commercial Proposal',
            'implementation': 'Implementation Plan',
            'architecture': 'Technical Architecture Document',
            'company': 'Company Profile & Credentials',
            'compliance': 'Compliance & Requirements Matrix'
        }
        return titles.get(deliverable_type, 'Proposal Document')

    def _get_document_description(self, deliverable_type: str) -> str:
        """Get document description"""
        descriptions = {
            'technical': 'Comprehensive technical solution and implementation approach',
            'commercial': 'Detailed commercial proposal with pricing and terms',
            'implementation': 'Project implementation plan with timeline and resources',
            'architecture': 'Technical architecture design and technology recommendations',
            'company': 'Company profile, credentials, and relevant experience',
            'compliance': 'Requirements traceability and compliance matrix'
        }
        return descriptions.get(deliverable_type, 'Proposal document')

    def _get_file_size(self, filename: str) -> str:
        """Get formatted file size"""
        try:
            filepath = os.path.join(self.output_dir, filename)
            size_bytes = os.path.getsize(filepath)

            if size_bytes < 1024:
                return f"{size_bytes} B"
            elif size_bytes < 1024 * 1024:
                return f"{size_bytes / 1024:.1f} KB"
            else:
                return f"{size_bytes / (1024 * 1024):.1f} MB"
        except:
            return "Unknown"

    # Fallback methods for when AI is not available
    def _fallback_technical_proposal(self) -> str:
        return f"""
# TECHNICAL PROPOSAL
## For {self.project.name}

### 1. EXECUTIVE SUMMARY
We are pleased to present our technical proposal for {self.project.name}. Our solution addresses all {len(self.analysis_results.get('must_have_requirements', []))} must-have requirements while providing a scalable, secure, and maintainable platform.

### 2. UNDERSTANDING OF REQUIREMENTS
Based on our analysis of the RFP documents, we understand that {self.client_name} requires:

{chr(10).join('• ' + req for req in self.analysis_results.get('must_have_requirements', [])[:10])}

### 3. PROPOSED SOLUTION ARCHITECTURE
Our solution follows industry best practices and modern architectural patterns to ensure scalability, security, and maintainability.

### 4. TECHNICAL IMPLEMENTATION APPROACH
We propose an agile development methodology with continuous integration and deployment practices.

### 5. COMPLIANCE & STANDARDS
Our solution complies with all relevant industry standards and regulatory requirements.

Note: AI analysis service not available. This is a basic template. Please configure ANTHROPIC_API_KEY for detailed AI-generated proposals.
"""

    def _fallback_commercial_proposal(self) -> str:
        return f"""
# COMMERCIAL PROPOSAL
## For {self.project.name}

### 1. EXECUTIVE SUMMARY
{self.company_name} is pleased to submit this commercial proposal for {self.project.name}.

### 2. PRICING STRUCTURE
Our competitive pricing reflects the value delivered and includes all necessary components for successful project delivery.

### 3. PROJECT TIMELINE
We propose a phased approach to ensure timely delivery and risk mitigation.

Note: AI analysis service not available. This is a basic template. Please configure ANTHROPIC_API_KEY for detailed AI-generated proposals.
"""

    def _fallback_implementation_plan(self) -> str:
        return f"""
# IMPLEMENTATION PLAN
## For {self.project.name}

### 1. PROJECT OVERVIEW
This implementation plan covers the delivery of {self.project.name} meeting all specified requirements.

### 2. PROJECT PHASES
- Phase 1: Planning & Design
- Phase 2: Development & Configuration
- Phase 3: Testing & Integration
- Phase 4: Deployment & Go-Live

Note: AI analysis service not available. This is a basic template. Please configure ANTHROPIC_API_KEY for detailed AI-generated proposals.
"""

    def _fallback_architecture_document(self) -> str:
        return f"""
# TECHNICAL ARCHITECTURE DOCUMENT
## For {self.project.name}

### 1. ARCHITECTURE OVERVIEW
This document outlines the technical architecture for {self.project.name}.

### 2. SYSTEM ARCHITECTURE
Our architecture follows modern best practices for scalability and maintainability.

Note: AI analysis service not available. This is a basic template. Please configure ANTHROPIC_API_KEY for detailed AI-generated proposals.
"""

    def _fallback_company_profile(self) -> str:
        return f"""
# COMPANY PROFILE & CREDENTIALS
## {self.company_name}

### 1. COMPANY OVERVIEW
{self.company_name} is a leading technology solutions provider with extensive experience in delivering complex projects.

### 2. RELEVANT EXPERIENCE
We have successfully delivered numerous projects similar to {self.project.name}.

Note: AI analysis service not available. This is a basic template. Please configure ANTHROPIC_API_KEY for detailed AI-generated proposals.
"""

    def _fallback_compliance_matrix(self) -> str:
        return f"""
# COMPLIANCE & REQUIREMENTS MATRIX
## For {self.project.name}

### 1. REQUIREMENTS COMPLIANCE
We commit to meeting all {len(self.analysis_results.get('must_have_requirements', []))} must-have requirements specified in the RFP.

{chr(10).join(f'• {req}' for req in self.analysis_results.get('must_have_requirements', []))}

Note: AI analysis service not available. This is a basic template. Please configure ANTHROPIC_API_KEY for detailed AI-generated proposals.
"""

    def _fallback_custom_document(self, custom_deliverable) -> str:
        """Generate fallback content for custom deliverable when AI is not available"""
        return f"""
# {custom_deliverable.title}
## For {self.project.name}

### 1. DOCUMENT OVERVIEW
This is a custom deliverable created for {self.project.name} by {self.company_name}.

### 2. PROJECT CONTEXT
- **Project:** {self.project.name}
- **Client:** {self.client_name}
- **Contact:** {self.contact_person}

### 3. CONTENT
{custom_deliverable.description}

Note: AI analysis service not available. This is a basic template. Please configure ANTHROPIC_API_KEY for detailed AI-generated custom deliverables using your prompt template:

**Your Template:**
{custom_deliverable.prompt_template or 'No template defined'}
"""

    def _format_requirements(self) -> str:
        """Format requirements for template replacement"""
        must_have = self.analysis_results.get('must_have_requirements', [])
        good_to_have = self.analysis_results.get('good_to_have_requirements', [])
        
        formatted = []
        if must_have:
            formatted.append("**Must-Have Requirements:**")
            formatted.extend([f"• {req}" for req in must_have])
        
        if good_to_have:
            formatted.append("\n**Good-to-Have Requirements:**")
            formatted.extend([f"• {req}" for req in good_to_have])
        
        return "\n".join(formatted) if formatted else "No specific requirements identified."

    def _format_analysis_summary(self) -> str:
        """Format analysis summary for template replacement"""
        summary_parts = []
        
        if self.analysis_results.get('must_have_requirements'):
            summary_parts.append(f"{len(self.analysis_results['must_have_requirements'])} must-have requirements identified")
        
        if self.analysis_results.get('technical_specifications'):
            summary_parts.append(f"{len(self.analysis_results['technical_specifications'])} technical specifications")
        
        if self.analysis_results.get('project_details'):
            details = self.analysis_results['project_details']
            if details.get('timeline'):
                summary_parts.append(f"Timeline: {details['timeline']}")
            if details.get('budget'):
                summary_parts.append(f"Budget: {details['budget']}")
        
        return "; ".join(summary_parts) if summary_parts else "Analysis completed for project requirements."

    def _format_constraints(self) -> str:
        """Format key constraints for template replacement"""
        constraints = []
        
        # Add any constraints from analysis results
        if self.analysis_results.get('constraints'):
            constraints.extend(self.analysis_results['constraints'])
        
        # Add timeline constraints
        project_details = self.analysis_results.get('project_details', {})
        if project_details.get('timeline'):
            constraints.append(f"Timeline constraint: {project_details['timeline']}")
        
        if project_details.get('budget'):
            constraints.append(f"Budget constraint: {project_details['budget']}")
        
        return "\n".join([f"• {constraint}" for constraint in constraints]) if constraints else "No specific constraints identified."

    def _format_project_details(self) -> str:
        """Format project details for template replacement"""
        details = self.analysis_results.get('project_details', {})
        formatted = []
        
        if details.get('timeline'):
            formatted.append(f"**Timeline:** {details['timeline']}")
        
        if details.get('budget'):
            formatted.append(f"**Budget:** {details['budget']}")
        
        if details.get('evaluation_criteria'):
            formatted.append(f"**Evaluation Criteria:** {details['evaluation_criteria']}")
        
        if details.get('key_stakeholders'):
            formatted.append(f"**Key Stakeholders:** {details['key_stakeholders']}")
        
        return "\n".join(formatted)

    def _get_past_proposal_insights(self, proposal_type: str) -> Dict[str, Any]:
        """Get enhanced past proposal insights using Claude Vector Intelligence"""
        try:
            from proposal_manager import get_proposal_manager
            
            proposal_manager = get_proposal_manager()
            insights = proposal_manager.get_relevant_past_proposals(
                analysis_results=self.analysis_results,
                proposal_type=proposal_type,
                limit=8
            )
            
            return insights
            
        except Exception as e:
            print(f"Warning: Could not retrieve past proposal insights: {e}")
            return {'success': False, 'error': str(e)}
    
    def _format_past_proposal_context(self, insights: Dict[str, Any]) -> str:
        """Format enhanced past proposal insights from Claude Vector Intelligence"""
        if not insights.get('success'):
            return ""
        
        context_parts = [
            "\n**🎯 CLAUDE VECTOR INTELLIGENCE - PAST PROPOSAL INSIGHTS:**",
            f"**Intelligence Score:** {insights.get('confidence_score', 0):.0%}",
            f"**Analyzed Sources:** {insights.get('found_proposals', 0)} relevant past proposals",
            f"**Generated:** {insights.get('intelligence_timestamp', 'Unknown')}"
        ]
        
        # Reusable content sections with enhanced intelligence
        reusable_content = insights.get('reusable_content', {})
        
        if reusable_content.get('technical_approach'):
            tech_approach = reusable_content['technical_approach']
            context_parts.extend([
                "\n**🔧 PROVEN TECHNICAL APPROACHES:**",
                f"Content: {tech_approach.get('content', '')[:400]}...",
                f"Adaptation Needed: {tech_approach.get('adaptation_needed', 'Standard customization')}",
                f"Confidence: {tech_approach.get('confidence', 0):.0%}",
                f"Source Proposals: {', '.join(tech_approach.get('source_proposals', []))}"
            ])
        
        if reusable_content.get('implementation_methodology'):
            impl_method = reusable_content['implementation_methodology']
            context_parts.extend([
                "\n**⚙️ PROVEN IMPLEMENTATION METHODOLOGIES:**",
                f"Content: {impl_method.get('content', '')[:400]}...",
                f"Adaptation Needed: {impl_method.get('adaptation_needed', 'Standard customization')}",
                f"Confidence: {impl_method.get('confidence', 0):.0%}"
            ])
        
        if reusable_content.get('team_experience'):
            team_exp = reusable_content['team_experience']
            context_parts.extend([
                "\n**👥 TEAM EXPERIENCE FROM PAST PROPOSALS:**",
                f"Content: {team_exp.get('content', '')[:400]}...",
                f"Adaptation Needed: {team_exp.get('adaptation_needed', 'Update for current project')}",
                f"Confidence: {team_exp.get('confidence', 0):.0%}"
            ])
        
        if reusable_content.get('solution_architecture'):
            sol_arch = reusable_content['solution_architecture']
            context_parts.extend([
                "\n**🏗️ SOLUTION ARCHITECTURE PATTERNS:**",
                f"Content: {sol_arch.get('content', '')[:400]}...",
                f"Adaptation Needed: {sol_arch.get('adaptation_needed', 'Customize for requirements')}",
                f"Confidence: {sol_arch.get('confidence', 0):.0%}"
            ])
        
        # Capability Intelligence
        capability_intel = insights.get('capability_intelligence', {})
        if capability_intel:
            context_parts.append("\n**💪 OUR PROVEN CAPABILITIES:**")
            if capability_intel.get('proven_capabilities'):
                context_parts.append(f"Capabilities: {', '.join(capability_intel['proven_capabilities'][:5])}")
            if capability_intel.get('technology_expertise'):
                context_parts.append(f"Technologies: {', '.join(capability_intel['technology_expertise'][:5])}")
            if capability_intel.get('competitive_differentiators'):
                context_parts.append(f"Differentiators: {', '.join(capability_intel['competitive_differentiators'][:3])}")
        
        # Gap Analysis
        gap_analysis = insights.get('gap_analysis', {})
        if gap_analysis:
            context_parts.append("\n**⚠️ GAP ANALYSIS:**")
            if gap_analysis.get('missing_capabilities'):
                context_parts.append(f"New Capabilities Needed: {', '.join(gap_analysis['missing_capabilities'][:3])}")
            if gap_analysis.get('research_needed'):
                context_parts.append(f"Research Required: {', '.join(gap_analysis['research_needed'][:3])}")
        
        # Generation Guidance
        generation_guidance = insights.get('generation_guidance', {})
        if generation_guidance:
            context_parts.append("\n**📝 GENERATION GUIDANCE:**")
            if generation_guidance.get('writing_style'):
                context_parts.append(f"Style: {generation_guidance['writing_style']}")
            if generation_guidance.get('key_messaging'):
                context_parts.append(f"Key Messages: {', '.join(generation_guidance['key_messaging'][:3])}")
            if generation_guidance.get('success_factors'):
                context_parts.append(f"Success Factors: {', '.join(generation_guidance['success_factors'][:3])}")
        
        context_parts.extend([
            "\n**🤖 CLAUDE GENERATION INSTRUCTIONS:**",
            "1. LEVERAGE proven technical approaches with appropriate adaptation",
            "2. REFERENCE similar past implementations where highly relevant", 
            "3. EMPHASIZE our proven capabilities and competitive differentiators",
            "4. ADDRESS any capability gaps with mitigation strategies",
            "5. MAINTAIN consistency with our established expertise and success patterns",
            "6. CUSTOMIZE all content to align perfectly with current RFP requirements",
            "7. USE the provided writing style and key messaging themes",
            "\n**🎯 INTELLIGENCE SYNTHESIS COMPLETE - USE FOR COMPETITIVE ADVANTAGE**\n"
        ])
        
        return "\n".join(context_parts)

    def _get_partner_solutions(self) -> Dict[str, Any]:
        """Get relevant partner solutions for current requirements"""
        try:
            # Import here to avoid circular imports
            import sys
            import os
            sys.path.append(os.path.dirname(os.path.abspath(__file__)))
            
            # Use the function from main.py
            from main import get_partner_solutions_for_requirements
            
            partner_solutions = get_partner_solutions_for_requirements(self.analysis_results)
            return partner_solutions
            
        except Exception as e:
            print(f"Warning: Could not retrieve partner solutions: {e}")
            return {'success': False, 'error': str(e)}

    def _format_partner_solutions_context(self, partner_solutions: Dict[str, Any]) -> str:
        """Format partner solutions for inclusion in AI prompt"""
        if not partner_solutions.get('success') or not partner_solutions.get('partner_solutions'):
            return ""
        
        solutions = partner_solutions.get('partner_solutions', [])
        context_parts = [
            "\n**RELEVANT PARTNER SOLUTIONS:**",
            f"**Evaluated Partners:** {partner_solutions.get('total_partners_evaluated', 0)}",
            f"**Recommended Solutions:** {len(solutions)}"
        ]
        
        for i, solution in enumerate(solutions[:3], 1):  # Top 3 partner solutions
            partner = solution['partner']
            matches = solution['matches']
            
            context_parts.append(f"\n**Partner Solution {i}: {partner['name']}**")
            context_parts.append(f"Website: {partner['website']}")
            context_parts.append(f"Capabilities: {partner['summary'][:200]}...")
            context_parts.append(f"Relevance Score: {matches['total_score']:.0%}")
            
            if matches['matching_offerings']:
                context_parts.append(f"Addresses: {', '.join(matches['matching_offerings'][:3])}")
            
            context_parts.append(f"Integration: {solution['recommended_integration']}")
            context_parts.append(f"Effort: {solution['estimated_effort']}")
        
        # Add guidance for partner integration
        context_parts.append("\n**Partner Integration Instructions:**")
        context_parts.append("1. Prioritize Temenos solutions as ITSS Global's strategic platform for core banking")
        context_parts.append("2. For requirements that ITSS Global/Temenos cannot fully address, recommend relevant vendor solutions")
        context_parts.append("3. Include vendor solutions in the 'Partner Solutions' section of the proposal")
        context_parts.append("4. Explain how vendor solutions integrate with core ITSS Global/Temenos offerings")
        context_parts.append("5. Highlight ITSS Global's BFSI expertise and Temenos partnership benefits")
        context_parts.append("6. Provide implementation timeline that includes vendor coordination")
        context_parts.append("")
        
        return "\n".join(context_parts)
