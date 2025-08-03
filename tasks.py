#!/usr/bin/env python3
"""
Updated tasks.py compatible with RealAnalysisEngine
"""
import os
import sys
from datetime import datetime
from celery import current_task
from main import celery, create_app

# Create app context for tasks
app = create_app()

def fix_document_schema():
    """Fix Document model schema issues"""
    print("🔧 Fixing database schema...")

    try:
        with app.app_context():
            from models import db

            engine = db.engine

            # Check current document table structure
            print("📊 Checking current schema...")

            try:
                inspector = db.inspect(engine)
                if 'documents' in inspector.get_table_names():
                    columns = inspector.get_columns('documents')
                    column_names = [col['name'] for col in columns]
                    print(f"   Current columns: {column_names}")

                    # Add missing columns if needed
                    missing_columns = []

                    if 'original_filename' not in column_names:
                        missing_columns.append('original_filename')
                    if 'processing_status' not in column_names:
                        missing_columns.append('processing_status')
                    if 'task_id' not in column_names:
                        missing_columns.append('task_id')
                    if 'error_message' not in column_names:
                        missing_columns.append('error_message')
                    if 'processed_at' not in column_names:
                        missing_columns.append('processed_at')
                    if 'content' not in column_names:  # For RealAnalysisEngine compatibility
                        missing_columns.append('content')

                    if missing_columns:
                        print(f"📝 Adding missing columns: {missing_columns}")

                        with engine.connect() as conn:
                            for col in missing_columns:
                                if col == 'original_filename':
                                    conn.execute(db.text("""
                                        ALTER TABLE documents
                                        ADD COLUMN IF NOT EXISTS original_filename VARCHAR(500);
                                    """))
                                    # Update existing records
                                    conn.execute(db.text("""
                                        UPDATE documents
                                        SET original_filename = filename
                                        WHERE original_filename IS NULL;
                                    """))
                                elif col == 'processing_status':
                                    conn.execute(db.text("""
                                        ALTER TABLE documents
                                        ADD COLUMN IF NOT EXISTS processing_status VARCHAR(50) DEFAULT 'pending';
                                    """))
                                elif col == 'task_id':
                                    conn.execute(db.text("""
                                        ALTER TABLE documents
                                        ADD COLUMN IF NOT EXISTS task_id VARCHAR(100);
                                    """))
                                elif col == 'error_message':
                                    conn.execute(db.text("""
                                        ALTER TABLE documents
                                        ADD COLUMN IF NOT EXISTS error_message TEXT;
                                    """))
                                elif col == 'processed_at':
                                    conn.execute(db.text("""
                                        ALTER TABLE documents
                                        ADD COLUMN IF NOT EXISTS processed_at TIMESTAMP;
                                    """))
                                elif col == 'content':
                                    conn.execute(db.text("""
                                        ALTER TABLE documents
                                        ADD COLUMN IF NOT EXISTS content TEXT;
                                    """))

                            conn.commit()
                            print("✅ Schema updated successfully")
                    else:
                        print("✅ Schema is already correct")

                    # Verify the fix
                    from models import Document
                    doc_count = Document.query.count()
                    print(f"✅ Verified: {doc_count} documents accessible")

                    return True
                else:
                    print("⚠️ Documents table not found")
                    return False

            except Exception as e:
                print(f"❌ Schema fix failed: {e}")
                return False

    except Exception as e:
        print(f"❌ Database connection failed: {e}")
        return False

@celery.task(bind=True, name='tasks.process_document_task')
def process_document_task(self, document_id):
    """Process uploaded document - compatible with RealAnalysisEngine"""
    with app.app_context():
        try:
            # Fix schema if needed
            fix_document_schema()

            current_task.update_state(state='PROGRESS', meta={'status': 'Starting document processing'})

            from models import db, Document
            document = Document.query.get(document_id)
            if not document:
                raise Exception(f"Document {document_id} not found")

            print(f"🤖 Processing document ID {document_id}")

            # Get filename safely
            filename = getattr(document, 'original_filename', None) or getattr(document, 'filename', 'unknown')
            print(f"📄 File: {filename}")

            # Update document status
            if hasattr(document, 'processing_status'):
                document.processing_status = 'in_progress'
                db.session.commit()

            current_task.update_state(state='PROGRESS', meta={'status': 'Extracting content from document'})

            # Extract content from file
            content = ""
            if document.file_path and os.path.exists(document.file_path):
                try:
                    file_ext = filename.lower().split('.')[-1] if '.' in filename else ''

                    if file_ext == 'txt':
                        # Try multiple encodings for text files
                        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
                        for encoding in encodings:
                            try:
                                with open(document.file_path, 'r', encoding=encoding, errors='ignore') as f:
                                    content = f.read()
                                    break
                            except UnicodeDecodeError:
                                continue

                    elif file_ext in ['docx', 'doc']:
                        try:
                            import docx
                            doc = docx.Document(document.file_path)
                            content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])

                            # Extract text from tables too
                            for table in doc.tables:
                                for row in table.rows:
                                    for cell in row.cells:
                                        content += ' ' + cell.text

                        except Exception as e:
                            print(f"⚠️ DOCX reading failed: {e}")
                            content = f"Document: {filename} (DOCX content extraction failed: {str(e)})"

                    elif file_ext == 'pdf':
                        try:
                            import PyPDF2
                            with open(document.file_path, 'rb') as file:
                                pdf_reader = PyPDF2.PdfReader(file)
                                content = ""
                                for page_num in range(len(pdf_reader.pages)):
                                    page = pdf_reader.pages[page_num]
                                    content += page.extract_text() + "\n\n"
                        except Exception as e:
                            print(f"⚠️ PDF reading failed: {e}")
                            content = f"Document: {filename} (PDF content extraction failed: {str(e)})"

                    elif file_ext in ['xlsx', 'xls']:
                        try:
                            import pandas as pd
                            xl_file = pd.ExcelFile(document.file_path)
                            content = ""
                            for sheet_name in xl_file.sheet_names:
                                df = pd.read_excel(document.file_path, sheet_name=sheet_name)
                                content += f"\n--- Sheet: {sheet_name} ---\n"
                                content += df.to_string(index=False) + "\n\n"
                        except Exception as e:
                            print(f"⚠️ Excel reading failed: {e}")
                            content = f"Document: {filename} (Excel content extraction failed: {str(e)})"
                    else:
                        content = f"Document: {filename} (unsupported format '{file_ext}' for text extraction)"

                    print(f"✅ Extracted {len(content)} characters")

                except Exception as e:
                    print(f"❌ File reading error: {e}")
                    content = f"Document: {filename} (file reading failed: {str(e)})"
            else:
                print(f"❌ File not found: {document.file_path}")
                content = f"Document: {filename} (file not found at {document.file_path})"

            current_task.update_state(state='PROGRESS', meta={'status': 'Storing extracted content'})

            # Store content in document for RealAnalysisEngine to use
            if hasattr(document, 'content'):
                document.content = content

            # Also store in uploaded_at if it doesn't exist
            if hasattr(document, 'uploaded_at') and not document.uploaded_at:
                document.uploaded_at = datetime.utcnow()

            db.session.commit()

            current_task.update_state(state='PROGRESS', meta={'status': 'Content extracted - ready for AI analysis'})

            # Mark as completed - RealAnalysisEngine will handle the AI analysis
            if hasattr(document, 'processing_status'):
                document.processing_status = 'completed'
            if hasattr(document, 'processed_at'):
                document.processed_at = datetime.utcnow()

            db.session.commit()

            print("💾 Document content extracted and stored - ready for RealAnalysisEngine")

            return {
                'status': 'completed',
                'document_id': document_id,
                'filename': filename,
                'content_length': len(content),
                'message': 'Document content extracted successfully - ready for AI analysis via RealAnalysisEngine'
            }

        except Exception as e:
            print(f"❌ Task failed: {e}")

            # Update document status on error
            try:
                document = Document.query.get(document_id)
                if document:
                    if hasattr(document, 'processing_status'):
                        document.processing_status = 'failed'
                    if hasattr(document, 'error_message'):
                        document.error_message = str(e)
                    db.session.commit()
            except:
                pass

            current_task.update_state(
                state='FAILURE',
                meta={'error': str(e), 'document_id': document_id}
            )
            raise

@celery.task(name='tasks.test_task')
def test_task():
    """Simple test task to verify Celery is working"""
    return "Celery is working correctly!"

@celery.task(name='tasks.cleanup_old_files')
def cleanup_old_files():
    """Cleanup old uploaded files"""
    with app.app_context():
        try:
            from models import Document
            import datetime

            # Find files older than 30 days
            cutoff_date = datetime.datetime.utcnow() - datetime.timedelta(days=30)
            old_docs = Document.query.filter(Document.uploaded_at < cutoff_date).all()

            cleaned_count = 0
            for doc in old_docs:
                try:
                    if doc.file_path and os.path.exists(doc.file_path):
                        os.remove(doc.file_path)
                        cleaned_count += 1
                except Exception as e:
                    print(f"Failed to remove file {doc.file_path}: {e}")

            return f"Cleanup completed: {cleaned_count} files removed"

        except Exception as e:
            return f"Cleanup failed: {str(e)}"

@celery.task(name='tasks.fix_schema_task')
def fix_schema_task():
    """Task to fix database schema"""
    with app.app_context():
        return fix_document_schema()

# Auto-run schema fix when tasks module is imported
if __name__ != '__main__':
    try:
        with app.app_context():
            fix_document_schema()
    except Exception as e:
        print(f"⚠️ Schema auto-fix failed: {e}")

if __name__ == '__main__':
    print("🚀 Running schema fix...")
    with app.app_context():
        result = fix_document_schema()
        print(f"Schema fix result: {result}")
