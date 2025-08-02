# agents/document_intelligence.py
import PyPDF2
import docx
import pytesseract
from PIL import Image
import json
import hashlib
import os
from typing import Dict, Any, List

from agents.base_agent import BaseAgent
from models import AgentTask, Document

class DocumentIntelligenceAgent(BaseAgent):
    """Specialized agent for document processing and analysis"""

    async def _execute_task(self, task: AgentTask) -> Dict[str, Any]:
        """Execute document intelligence tasks"""

        task_type = task.task_type
        input_data = task.input_data

        if task_type == 'document_analysis':
            return await self._analyze_document(input_data['document_id'])
        elif task_type == 'text_extraction':
            return await self._extract_text(input_data['document_id'])
        elif task_type == 'structure_analysis':
            return await self._analyze_structure(input_data['document_id'])
        else:
            raise ValueError(f"Unknown task type: {task_type}")

    async def _analyze_document(self, document_id: int) -> Dict[str, Any]:
        """Comprehensive document analysis"""

        document = Document.query.get(document_id)
        if not document:
            raise ValueError(f"Document {document_id} not found")

        # Extract text content
        extracted_text = await self._extract_text_from_file(document.file_path, document.mime_type)

        # Analyze document structure using Claude
        structure_prompt = f"""
        Analyze this document and extract its structure:

        Document Title: {document.original_filename}
        Content: {extracted_text[:5000]}...

        Please identify:
        1. Document type (RFP, technical specification, legal document, etc.)
        2. Main sections and their hierarchy
        3. Key information blocks (requirements, deadlines, contact info)
        4. Tables and structured data
        5. Important entities (companies, dates, amounts, technical terms)

        Return the analysis as structured JSON.
        """

        structure_analysis = await self.call_claude(structure_prompt)

        # Update document in database
        document.extracted_text = extracted_text
        document.extracted_metadata = json.loads(structure_analysis)
        document.processing_status = 'completed'
        db.session.commit()

        return {
            'document_id': document_id,
            'extracted_text': extracted_text,
            'structure_analysis': structure_analysis,
            'metadata': document.extracted_metadata,
            'word_count': len(extracted_text.split()),
            'status': 'completed'
        }

    async def _extract_text_from_file(self, file_path: str, mime_type: str) -> str:
        """Extract text from various file formats"""

        if mime_type == 'application/pdf':
            return self._extract_from_pdf(file_path)
        elif mime_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
            return self._extract_from_word(file_path)
        else:
            # Try OCR for image files
            return self._extract_with_ocr(file_path)

    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF"""

        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            self._log_event('ERROR', 'pdf_extraction_failed', f"Failed to extract PDF: {str(e)}")
            # Fallback to OCR
            text = self._extract_with_ocr(file_path)

        return text

    def _extract_from_word(self, file_path: str) -> str:
        """Extract text from Word documents"""

        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"

            return text

        except Exception as e:
            self._log_event('ERROR', 'word_extraction_failed', f"Failed to extract Word: {str(e)}")
            return ""

    def _extract_with_ocr(self, file_path: str) -> str:
        """Extract text using OCR for scanned documents"""

        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            return text
        except Exception as e:
            self._log_event('ERROR', 'ocr_failed', f"OCR extraction failed: {str(e)}")
            return ""
